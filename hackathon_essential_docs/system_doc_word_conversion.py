"""
GoalFlow — Hackathon Submission Word Document Generator
========================================================

Reads hackathon_essential_docs/GoalFlow_Submission.md and produces a
polished `.docx` ready for upload to Unstop (AtomQuest Hackathon 2026).

- Embeds the Atomberg logo header
- Renders every PlantUML ```plantuml ... ``` block as a PNG (via the
  public PlantUML server) and inserts it inline
- Embeds every screenshot referenced as ![](../screenshots/xxx.png)
- Preserves headings, tables, bullet/numbered lists, bold/italic and
  hyperlinks

Run:
    pip install python-docx requests markdown
    python system_doc_word_conversion.py

Output:
    GoalFlow_Submission.docx  (next to this script)
"""

from __future__ import annotations

import base64
import os
import re
import sys
import zlib
from pathlib import Path
from typing import Iterable
from urllib.parse import urlparse

import requests
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "hackathon_essential_docs" / "GoalFlow_Submission.md"
DOCX_PATH = ROOT / "GoalFlow_Submission.docx"
LOGO_PATH = ROOT / "apps" / "web" / "public" / "atomquest_log.png"
DIAGRAM_CACHE = ROOT / ".plantuml_cache"
DIAGRAM_CACHE.mkdir(exist_ok=True)


# ---------------------------------------------------------------------------
# PlantUML rendering via public server
# (Uses the standard deflate + custom base64 encoding)
# ---------------------------------------------------------------------------

_PLANTUML_ALPHABET = (
    "0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz-_"
)


def _plantuml_encode(text: str) -> str:
    """Encode PlantUML source to the URL-safe form used by plantuml.com."""
    data = text.encode("utf-8")
    compressed = zlib.compress(data)[2:-4]  # strip zlib header / checksum

    res = []
    i = 0
    while i < len(compressed):
        b1 = compressed[i]
        b2 = compressed[i + 1] if i + 1 < len(compressed) else 0
        b3 = compressed[i + 2] if i + 2 < len(compressed) else 0
        c1 = b1 >> 2
        c2 = ((b1 & 0x3) << 4) | (b2 >> 4)
        c3 = ((b2 & 0xF) << 2) | (b3 >> 6)
        c4 = b3 & 0x3F
        res.extend(
            [
                _PLANTUML_ALPHABET[c1 & 0x3F],
                _PLANTUML_ALPHABET[c2 & 0x3F],
                _PLANTUML_ALPHABET[c3 & 0x3F],
                _PLANTUML_ALPHABET[c4 & 0x3F],
            ]
        )
        i += 3
    return "".join(res)


def render_plantuml(source: str, name_hint: str) -> Path | None:
    """Fetch a PNG render of the given PlantUML source. Cached on disk."""
    safe_name = re.sub(r"[^A-Za-z0-9_-]+", "_", name_hint).strip("_") or "diagram"
    cache_file = DIAGRAM_CACHE / f"{safe_name}.png"
    if cache_file.exists() and cache_file.stat().st_size > 0:
        return cache_file

    try:
        encoded = _plantuml_encode(source)
        url = f"https://www.plantuml.com/plantuml/png/{encoded}"
        print(f"  ↳ rendering PlantUML: {safe_name} ({url[:80]}...)")
        resp = requests.get(url, timeout=30)
        resp.raise_for_status()
        cache_file.write_bytes(resp.content)
        return cache_file
    except Exception as exc:  # noqa: BLE001
        print(f"  ⚠️  PlantUML render failed for {safe_name}: {exc}")
        return None


# ---------------------------------------------------------------------------
# Lightweight Markdown → docx writer
# (Tailored to the structure of GoalFlow_Submission.md — not a full MD parser)
# ---------------------------------------------------------------------------

INLINE_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
INLINE_BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
INLINE_ITALIC_RE = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")
IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Add a clickable hyperlink to a python-docx paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F6FEB")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _emit_inline(paragraph, text: str, bold: bool = False) -> None:
    """Emit a span of text with inline markdown (bold, italic, code, link)."""
    # Strip leading/trailing whitespace artifacts but preserve internal spaces
    if not text:
        return

    # We tokenize iteratively against the first match of any inline pattern
    patterns = [
        ("link", INLINE_LINK_RE),
        ("bold", INLINE_BOLD_RE),
        ("italic", INLINE_ITALIC_RE),
        ("code", INLINE_CODE_RE),
    ]

    pos = 0
    while pos < len(text):
        best = None
        for kind, regex in patterns:
            m = regex.search(text, pos)
            if m and (best is None or m.start() < best[1].start()):
                best = (kind, m)

        if best is None:
            run = paragraph.add_run(text[pos:])
            if bold:
                run.bold = True
            return

        kind, m = best
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            if bold:
                run.bold = True

        if kind == "link":
            _add_hyperlink(paragraph, m.group(2), m.group(1))
        elif kind == "bold":
            run = paragraph.add_run(m.group(1))
            run.bold = True
        elif kind == "italic":
            run = paragraph.add_run(m.group(1))
            run.italic = True
        elif kind == "code":
            run = paragraph.add_run(m.group(1))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)

        pos = m.end()


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            _emit_inline(p, cell_text.strip(), bold=(r_idx == 0))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx == 0:
                _shade_cell(cell, "1F6FEB")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True
    doc.add_paragraph()


def _add_code_block(doc: Document, code: str, language: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Shade paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F5F7")
    pPr.append(shd)

    run = p.add_run(code.rstrip("\n"))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x24, 0x29, 0x2F)


_IMAGE_MAGIC = {
    b"\x89PNG\r\n\x1a\n": "png",
    b"\xff\xd8\xff": "jpg",
    b"GIF87a": "gif",
    b"GIF89a": "gif",
    b"BM": "bmp",
}


def _is_real_image(path: Path) -> bool:
    try:
        head = path.read_bytes()[:16]
    except OSError:
        return False
    return any(head.startswith(sig) for sig in _IMAGE_MAGIC)


def _add_image(doc: Document, image_path: Path, width_inches: float = 6.0) -> None:
    if not image_path.exists():
        print(f"  ⚠️  missing image: {image_path}")
        p = doc.add_paragraph()
        run = p.add_run(f"[Missing image: {image_path.name}]")
        run.italic = True
        return
    if not _is_real_image(image_path):
        print(f"  ⚠️  not a valid image (skipping): {image_path}")
        p = doc.add_paragraph()
        run = p.add_run(f"[Skipped non-image file: {image_path.name}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    except Exception as exc:  # noqa: BLE001
        print(f"  ⚠️  failed to embed {image_path.name}: {exc}")
        run = p.add_run(f"[Could not embed {image_path.name}: {exc}]")
        run.italic = True


def _add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x57, 0x60, 0x6A)


# ---------------------------------------------------------------------------
# Markdown line-level state machine
# ---------------------------------------------------------------------------


def _resolve_path(rel: str) -> Path:
    """Resolve a path relative to the markdown file."""
    if rel.startswith(("http://", "https://")):
        return Path(rel)  # caller checks .exists() — won't, so handled inline
    md_dir = MD_PATH.parent
    return (md_dir / rel).resolve()


def _strip_html(text: str) -> str:
    """Drop trivial HTML wrappers we use in the MD header for centering."""
    return re.sub(r"</?(div|center)[^>]*>", "", text, flags=re.IGNORECASE).strip()


def _is_table_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*", line))


def _split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def convert(md_text: str, doc: Document) -> None:
    lines = md_text.splitlines()
    i = 0
    diagram_count = 0
    screenshot_count = 0

    while i < len(lines):
        line = lines[i]

        # --- Fenced code block --------------------------------------------------
        if line.lstrip().startswith("```"):
            lang = line.lstrip()[3:].strip().lower()
            buf: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].lstrip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # consume closing fence
            code = "\n".join(buf)

            if lang == "plantuml":
                diagram_count += 1
                name_match = re.search(r"@startuml\s+(\S+)", code)
                name_hint = name_match.group(1) if name_match else f"diagram_{diagram_count}"
                png = render_plantuml(code, name_hint)
                if png:
                    _add_image(doc, png, width_inches=6.2)
                    _add_caption(doc, f"Figure {diagram_count}: {name_hint.replace('-', ' ').replace('_', ' ')}")
                else:
                    _add_code_block(doc, code, lang)
            else:
                _add_code_block(doc, code, lang)
            continue

        # --- Image standalone line ---------------------------------------------
        img_match = IMAGE_RE.fullmatch(line.strip()) if line.strip() else None
        if img_match:
            alt, src = img_match.group(1), img_match.group(2)
            path = _resolve_path(src)
            screenshot_count += 1
            _add_image(doc, path, width_inches=6.2)
            if alt:
                _add_caption(doc, f"Screenshot {screenshot_count}: {alt}")
            i += 1
            continue

        # --- Horizontal rule ---------------------------------------------------
        if line.strip() in {"---", "***", "___"}:
            doc.add_paragraph().add_run().add_break()
            i += 1
            continue

        # --- Headings ----------------------------------------------------------
        h_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h_match:
            level = len(h_match.group(1))
            text = _strip_html(h_match.group(2))
            heading = doc.add_heading(level=min(level, 4))
            heading.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
            )
            _emit_inline(heading, text)
            i += 1
            continue

        # --- Tables ------------------------------------------------------------
        if "|" in line and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            header = _split_table_row(line)
            i += 2  # skip header and separator
            body: list[list[str]] = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                body.append(_split_table_row(lines[i]))
                i += 1
            _add_table(doc, [header] + body)
            continue

        # --- Blockquote --------------------------------------------------------
        if line.lstrip().startswith(">"):
            buf = []
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                buf.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "EAF3FF")
            pPr.append(shd)
            _emit_inline(p, " ".join(s.strip() for s in buf if s.strip()))
            continue

        # --- Bullet list -------------------------------------------------------
        bullet_match = re.match(r"^\s*[-*+]\s+(.*)$", line)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            _emit_inline(p, bullet_match.group(1))
            i += 1
            continue

        # --- Numbered list -----------------------------------------------------
        num_match = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if num_match:
            p = doc.add_paragraph(style="List Number")
            _emit_inline(p, num_match.group(1))
            i += 1
            continue

        # --- Blank line --------------------------------------------------------
        if not line.strip():
            i += 1
            continue

        # --- Plain paragraph (strip stray HTML used for layout) ----------------
        text = _strip_html(line)
        if text:
            # Inline image inside a paragraph? Replace with embed.
            img_inline = IMAGE_RE.search(text)
            if img_inline and not text.strip().replace(img_inline.group(0), "").strip():
                # Pure image with leftover whitespace
                path = _resolve_path(img_inline.group(2))
                _add_image(doc, path, width_inches=4.0)
                i += 1
                continue
            # HTML <img> used in the banner block
            html_img = re.search(r'<img[^>]+src="([^"]+)"[^>]*?(?:width="?(\d+)"?)?', text)
            if html_img:
                src = html_img.group(1)
                width_attr = html_img.group(2)
                path = _resolve_path(src)
                width_in = float(width_attr) / 96.0 if width_attr else 2.2
                _add_image(doc, path, width_inches=max(width_in, 1.5))
                i += 1
                continue
            p = doc.add_paragraph()
            _emit_inline(p, text)
        i += 1


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------


def _configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for lvl, size, color in [
        (1, 24, RGBColor(0x0B, 0x4F, 0xA6)),
        (2, 18, RGBColor(0x1F, 0x6F, 0xEB)),
        (3, 14, RGBColor(0x24, 0x29, 0x2F)),
        (4, 12, RGBColor(0x57, 0x60, 0x6A)),
    ]:
        try:
            s = styles[f"Heading {lvl}"]
            s.font.size = Pt(size)
            s.font.color.rgb = color
            s.font.name = "Calibri"
        except KeyError:
            pass


def _ensure_logo() -> Path | None:
    """Return a valid local PNG of the cover logo."""
    if LOGO_PATH.exists() and _is_real_image(LOGO_PATH):
        return LOGO_PATH
    print(f"  ⚠️  Logo not found or invalid at {LOGO_PATH}")
    return None


def _add_cover(doc: Document) -> None:
    logo = _ensure_logo()
    if logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(str(logo), width=Inches(2.4))
        except Exception as exc:  # noqa: BLE001
            print(f"  ⚠️  Cover logo embed failed: {exc}")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GoalFlow")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x0B, 0x4F, 0xA6)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Enterprise Goal Setting & Tracking Portal")
    run.italic = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x57, 0x60, 0x6A)

    badge = doc.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = badge.add_run("AtomQuest Hackathon 2026 — Submission via Unstop")
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()  # spacer

    meta = [
        ("Candidate", "Gurjas Singh Gandhi"),
        ("College", "P.E.S's Modern College of Engineering, Pune"),
        ("Course / Year", "MCA, 2026"),
        ("Live Application", "https://goal-flow-theta.vercel.app"),
        ("Source Code (GitHub)", "https://github.com/Gurjas2112/GoalFlow"),
        ("Submission Date", "May 19, 2026"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Light List Accent 1"
    for r, (k, v) in enumerate(meta):
        c0 = table.rows[r].cells[0]
        c1 = table.rows[r].cells[1]
        c0.text = ""
        c1.text = ""
        run = c0.paragraphs[0].add_run(k)
        run.bold = True
        if v.startswith("http"):
            _add_hyperlink(c1.paragraphs[0], v, v)
        else:
            c1.paragraphs[0].add_run(v)
    doc.add_page_break()


def _set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> int:
    if not MD_PATH.exists():
        print(f"❌ Source markdown not found: {MD_PATH}")
        return 1

    print(f"📄 Reading  : {MD_PATH.relative_to(ROOT)}")
    md_text = MD_PATH.read_text(encoding="utf-8")

    # Strip the original front-matter cover block (between the first banner
    # <div align="center"> ... </div> and the first "---" thematic break).
    # Our generated cover page replaces it.
    md_after_cover = re.sub(
        r"\A<div align=\"center\">.*?</div>\s*---\s*",
        "",
        md_text,
        count=1,
        flags=re.DOTALL,
    )

    doc = Document()
    _set_margins(doc)
    _configure_styles(doc)
    _add_cover(doc)

    convert(md_after_cover, doc)

    doc.save(DOCX_PATH)
    print(f"✅ Wrote    : {DOCX_PATH.relative_to(ROOT)} ({DOCX_PATH.stat().st_size // 1024} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
